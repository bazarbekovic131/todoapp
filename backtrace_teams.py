
import sys
from PyQt5.QtWidgets import QApplication, QWidget,QComboBox, QVBoxLayout, QFileDialog, QLabel, QPushButton, QLineEdit, QMessageBox
from PyQt5.QtGui import QIcon

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from openpyxl import Workbook

def export_to_pdf(data, filename="schedule.pdf"):
    c = canvas.Canvas(f'{filename}.pdf', pagesize=letter)
    c.drawString(100, 750, "Schedule")
    for idx, (group, activities) in enumerate(data.items(), start=1):
        c.drawString(100, 750 - 15*idx, f"Group {group}: {activities[0]}, {activities[1]}")
    c.save()

def export_to_word(data, filename="schedule.docx"):
    doc = Document()
    doc.add_heading('Schedule', 0)
    for group, activities in data.items():
        doc.add_paragraph(f"Group {group}: {activities[0]}, {activities[1]}")
    doc.save(filename)

def export_to_excel(data, filename="schedule.xlsx"):
    wb = Workbook()
    ws = wb.active
    ws['A1'] = 'Group'
    ws['B1'] = 'Activity 1'
    ws['C1'] = 'Activity 2'
    for idx, (group, activities) in enumerate(data.items(), start=2):
        ws[f'A{idx}'] = f"Group {group}"
        ws[f'B{idx}'] = activities[0]
        ws[f'C{idx}'] = activities[1]
    wb.save(filename)

def is_valid(assignment, current_group_id, activity1, activity2):
    for group_id, activities in assignment.items():
        if group_id != current_group_id:
            if activity1 in activities or activity2 in activities:
                return False
    return True

def backtrack(group_id, assignment, activities, num_groups):
    if group_id > num_groups:
        return True, assignment
    for activity1 in activities:
        for activity2 in activities:
            if activity1 != activity2 and is_valid(assignment, group_id, activity1, activity2):
                assignment[group_id] = [activity1, activity2]
                if backtrack(group_id + 1, assignment, activities, num_groups)[0]:
                    return True, assignment
                assignment[group_id] = []
    return False, {}

class SchedulerApp(QWidget):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Activity Scheduler')
        self.setGeometry(100, 100, 400, 200)
        layout = QVBoxLayout()

        self.numGroupsInput = QLineEdit(self)
        self.numGroupsInput.setPlaceholderText('Enter number of groups (1-10)')
        layout.addWidget(self.numGroupsInput)

        self.formatComboBox = QComboBox(self)
        self.formatComboBox.addItems(['PDF', 'Word', 'Excel'])
        layout.addWidget(self.formatComboBox)

        self.fileNameInput = QLineEdit(self)
        self.fileNameInput.setPlaceholderText('Enter or select file name')
        layout.addWidget(self.fileNameInput)

        self.browseButton = QPushButton('Browse', self)
        self.browseButton.clicked.connect(self.openFileDialog)
        layout.addWidget(self.browseButton)

        self.saveButton = QPushButton('Save Schedule', self)
        self.saveButton.clicked.connect(self.save_schedule)
        layout.addWidget(self.saveButton)

        self.scheduleButton = QPushButton('Generate Schedule', self)
        self.scheduleButton.clicked.connect(self.generate_schedule)
        self.scheduleButton.setIcon(QIcon("schedule.png"))
        layout.addWidget(self.scheduleButton)

        self.resultLabel = QLabel('Schedule will appear here.', self)
        layout.addWidget(self.resultLabel)
        self.setStyleSheet("""
            QWidget {
                font-size: 16px;
            }
            QPushButton {
                background-color: #006dbf;
                color: white;
                border-radius: 5px;
                padding: 10px;
            }
            QPushButton:hover {
                background-color: #005fa1;
            }
            QLineEdit {
                border: 2px solid #ccc;
                border-radius: 4px;
                padding: 5px;
            }
        """)

        self.setLayout(layout)

    def generate_schedule(self):
        num_groups = int(self.numGroupsInput.text())
        activities = ["Football", "Basketball", "Hockey", "Tennis", "Volleyball", 
                      "Ultimate", "Squash", "Lacrosse", "American Football", "Softball"]
        assignment = {i: [] for i in range(1, num_groups + 1)}
        result, final_assignment = backtrack(1, assignment, activities, num_groups)
        if result:
            self.final_assignment = final_assignment
            schedule_text = "\n".join([f"Group {group}: {acts[0]}, {acts[1]}" for group, acts in final_assignment.items()])
            self.resultLabel.setText(schedule_text)
        else:
            QMessageBox.information(self, "Result", "No solution could be found.", QMessageBox.Ok)

    def save_schedule(self):
        if hasattr(self, 'final_assignment'):
            filename = self.fileNameInput.text()
            format_choice = self.formatComboBox.currentText()
            if format_choice == 'PDF':
                export_to_pdf(self.final_assignment, filename)
            elif format_choice == 'Word':
                export_to_word(self.final_assignment, filename)
            elif format_choice == 'Excel':
                export_to_excel(self.final_assignment, filename)
            QMessageBox.information(self, "Export", f"Schedule exported as {format_choice}.", QMessageBox.Ok)
        else:
            QMessageBox.warning(self, "Export", "Generate schedule before saving.", QMessageBox.Ok)

    def openFileDialog(self):
        options = QFileDialog.Options()
        fileName, _ = QFileDialog.getSaveFileName(self, "QFileDialog.getSaveFileName()", "", "All Files (*);;PDF Files (*.pdf);;Word Files (*.docx);;Excel Files (*.xlsx)", options=options)
        if fileName:
            self.fileNameInput.setText(fileName)

if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = SchedulerApp()
    ex.show()
    sys.exit(app.exec_())
